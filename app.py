import streamlit as st
import os
import time
from typing import List, Tuple
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import SentenceTransformerEmbeddings
from langchain.vectorstores import FAISS
from langchain.schema import Document as LC_Document
import google.generativeai as genai

# -----------------------------
# Configuration / Utilities
# -----------------------------

def configure_gemini():
    """
    Configures the Google Generative AI (Gemini) client using Streamlit secrets.
    Returns a configured gemini model object or None on failure.
    """
    try:
        google_api_key = st.secrets.get("GOOGLE_API_KEY")
        if not google_api_key:
            st.error("Google API Key not found. Please add it to your Streamlit secrets.")
            return None
        genai.configure(api_key=google_api_key)
        model = genai.GenerativeModel("gemini-2.5-flash")
        return model
    except Exception as e:
        st.error(f"Failed to configure Gemini: {e}")
        return None

# -----------------------------
# File reading & splitting (produce LangChain Documents with metadata)
# -----------------------------

def get_text_from_files(uploaded_files) -> List[LC_Document]:
    """
    Returns a list of LangChain Document objects with metadata:
      metadata['source'] = filename
      metadata['page']   = page number (1-indexed) or paragraph index for docx
    """
    docs: List[LC_Document] = []
    for file in uploaded_files:
        try:
            fname = file.name
            if fname.lower().endswith('.pdf'):
                pdf_reader = PdfReader(file)
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        docs.append(LC_Document(
                            page_content=page_text,
                            metadata={'source': fname, 'page': i + 1}
                        ))
            elif fname.lower().endswith('.docx'):
                doc = DocxDocument(file)
                # docx doesn't give page numbers reliably; use paragraph index as 'page'
                for i, para in enumerate(doc.paragraphs):
                    text = (para.text or "").strip()
                    if text:
                        docs.append(LC_Document(
                            page_content=text,
                            metadata={'source': fname, 'page': i + 1}
                        ))
            elif fname.lower().endswith('.txt'):
                txt = file.getvalue().decode("utf-8")
                if txt.strip():
                    docs.append(LC_Document(
                        page_content=txt,
                        metadata={'source': fname, 'page': 1}
                    ))
        except Exception as e:
            st.error(f"Error reading {file.name}: {e}")
    return docs


def get_text_chunks(documents: List[LC_Document], chunk_size=1000, chunk_overlap=200) -> List[LC_Document]:
    """Split LangChain Documents into chunks while preserving metadata."""
    if not documents:
        return []
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=[" ", " ", " "]
    )
    # split_documents preserves metadata on each chunk
    chunks = splitter.split_documents(documents)
    return chunks

# -----------------------------
# Vector store creation using LangChain
# -----------------------------

def create_vector_store(doc_chunks: List[LC_Document]):
    """Creates a LangChain FAISS vectorstore from Document chunks (with metadata)."""
    if not doc_chunks:
        st.warning("No text to process. Please upload and process valid documents.")
        return None
    try:
        embeddings = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
        # use from_documents so we keep metadata with each vector
        vectorstore = FAISS.from_documents(documents=doc_chunks, embedding=embeddings)
        return vectorstore
    except Exception as e:
        st.error(f"Failed to create vector store: {e}")
        return None

# -----------------------------
# Querying / Response generation with streaming display
# -----------------------------

def get_gemini_response(gemini_model, user_question: str, vectorstore) -> Tuple[str, List[Tuple[str, int]]]:
    """
    Retrieves top-k relevant document chunks (each chunk has metadata with source & page),
    adds that context to the prompt and also returns a list of source citations for UI display.
    Returns tuple: (response_text, sources_list)
    """
    if not user_question:
        return "", []

    # 1. Retrieve top-k relevant documents using LangChain's similarity search
    try:
        k = 5
        docs = vectorstore.similarity_search(user_question, k=k)
        # Build context with inline source markers
        context_parts = []
        sources: List[Tuple[str, int]] = []  # unique (filename, page) pairs for UI
        for d in docs:
            src = d.metadata.get('source', 'unknown')
            page = d.metadata.get('page', 'unknown')
            context_parts.append(f"{d.page_content}\n\n[SOURCE: {src} - Page {page}]\n")
            pair = (src, page)
            if pair not in sources:
                sources.append(pair)
        context = "\n---\n".join(context_parts)
    except Exception as e:
        st.error(f"Failed retrieving relevant documents: {e}")
        context = ""
        sources = []

    # 2. Build the prompt (explicitly ask for citation of doc+page if used)
    chat_history_str = " ".join([f"{msg['role']}: {msg['content']}" for msg in st.session_state.get('chat_history', [])])

    prompt = f"""
You are DocuMind, an intelligent assistant. Use the provided context from the documents to answer the user's question.
When you use content from the context, include an inline citation after the sentence in this format: (source_filename - page N).
If the answer is not in the documents, explicitly state: "This answer is from my general knowledge base as I couldn't find the information in the provided document(s)."

CONTEXT:
{context}

CHAT HISTORY:
{chat_history_str}

USER QUESTION:
{user_question}

Helpful Answer (include citations like: (file.pdf - page 3) where relevant):
"""

    # 3. Call the Gemini API
    try:
        if gemini_model is None:
            return "Gemini model is not configured.", sources
        response = gemini_model.generate_content(prompt)
        return response.text, sources
    except Exception as e:
        st.error(f"Error generating response from Gemini: {e}")
        return "Sorry, I encountered an error while generating a response.", sources

# -----------------------------
# Streamlit App
# -----------------------------

def main():
    st.set_page_config(page_title="DocuMind (LangChain)", page_icon="🧠", layout="wide")

    gemini_model = configure_gemini()

    # Session state
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "vector_store_data" not in st.session_state:
        st.session_state.vector_store_data = {'vectorstore': None}

    # Sidebar
    with st.sidebar:
        st.title("DocuMind 🧠")
        st.markdown("Your document assistant using **LLM**. Upload your files and ask questions.")

        uploaded_files = st.file_uploader(
            "Upload your Documents (PDF, DOCX, TXT)",
            accept_multiple_files=True,
            type=['pdf', 'docx', 'txt']
        )

        if st.button("Process Documents", use_container_width=True, type="primary"):
            if uploaded_files:
                with st.spinner("Reading files..."):
                    docs = get_text_from_files(uploaded_files)
                with st.spinner("Chunking text..."):
                    text_chunks = get_text_chunks(docs)
                with st.spinner("Creating vector store... This may take a moment."):
                    vectorstore = create_vector_store(text_chunks)
                    if vectorstore is not None:
                        st.session_state.vector_store_data['vectorstore'] = vectorstore
                        st.session_state.chat_history = []
                        st.success("Documents processed successfully!")
            else:
                st.warning("Please upload at least one document.")

        # Clear chat option
        if st.button("Clear Chat", use_container_width=True):
            st.session_state.chat_history = []
            st.success("Chat cleared.")

    # Main chat
    st.header("Ask DocuMind Anything")

    if not st.session_state.vector_store_data['vectorstore']:
        st.info("Welcome! Please upload and process your documents in the sidebar to begin.")

    # Render chat history
    for message in st.session_state.chat_history:
        with st.chat_message(message["role"], avatar="👤" if message["role"] == "user" else "🤖"):
            st.markdown(message["content"])

    user_question = st.chat_input("Ask a question about your documents...")
    if user_question:
        if st.session_state.vector_store_data['vectorstore'] and gemini_model:
            # Save user message and render immediately
            st.session_state.chat_history.append({"role": "user", "content": user_question})
            with st.chat_message("user", avatar="👤"):
                st.markdown(user_question)

            # Prepare assistant message area and stream text into it
            with st.chat_message("assistant", avatar="🤖"):
                placeholder = st.empty()
                # Get full response text (we will display it incrementally)
                full_response, sources = get_gemini_response(
                    gemini_model=gemini_model,
                    user_question=user_question,
                    vectorstore=st.session_state.vector_store_data['vectorstore']
                )

                # If an error string returned, show it immediately
                if not full_response:
                    placeholder.markdown("(no response)")
                    st.session_state.chat_history.append({"role": "assistant", "content": "(no response)"})
                else:
                    # Stream by small chunks to simulate token-by-token / typing behavior
                    display_text = ""
                    chunk_size = 120  # characters per update
                    for i in range(0, len(full_response), chunk_size):
                        display_text += full_response[i:i+chunk_size]
                        # update placeholder
                        placeholder.markdown(display_text)
                        # small pause to create a typing effect
                        time.sleep(0.1)
                    # ensure final text is shown
                    placeholder.markdown(display_text)
                    # Save assistant response to history
                    st.session_state.chat_history.append({"role": "assistant", "content": full_response})

                    # After saving, show the sources used (if any)
                    if sources:
                        src_md = "### Sources referenced:\n"
                        for src, pg in sources:
                            src_md += f"- **{src}** — page {pg}\n"
                        st.markdown(src_md)
        else:
            st.warning("Please process your documents first or ensure the API key is configured correctly.")

if __name__ == "__main__":
    main()


